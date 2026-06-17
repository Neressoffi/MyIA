"""Generation de documents telechargeables a partir de texte Markdown.

JARVIS redige un contenu en Markdown simple (titres #, listes -, gras **...**),
et ce module le transforme en fichier : .txt, .md, .docx (Word) ou .pdf.
"""

from __future__ import annotations

import io
import re
import unicodedata

# Formats proposes : extension -> (libelle, type MIME).
FORMATS = {
    "txt": ("Texte brut", "text/plain; charset=utf-8"),
    "md": ("Markdown", "text/markdown; charset=utf-8"),
    "docx": (
        "Word (.docx)",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    "pdf": ("PDF", "application/pdf"),
}


# ----------------------------------------------------------------------
# Outils communs
# ----------------------------------------------------------------------
def nom_fichier(titre: str, fmt: str) -> str:
    """Construit un nom de fichier propre a partir d'un titre libre."""
    base = (titre or "document").strip().lower()
    base = unicodedata.normalize("NFKD", base).encode("ascii", "ignore").decode()
    base = re.sub(r"[^a-z0-9]+", "_", base).strip("_")
    base = base[:50] or "document"
    ext = fmt if fmt in FORMATS else "txt"
    return f"{base}.{ext}"


def _segments_gras(ligne: str):
    """Decoupe une ligne en segments (texte, gras?) sur les **...**."""
    segments = []
    for i, morceau in enumerate(re.split(r"\*\*(.+?)\*\*", ligne)):
        if morceau:
            segments.append((morceau, i % 2 == 1))
    return segments or [(ligne, False)]


def _markdown_simple(texte: str) -> str:
    """Retire le balisage Markdown pour un rendu texte brut lisible."""
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", texte)
    t = re.sub(r"`{1,3}", "", t)
    lignes = []
    for ligne in t.split("\n"):
        m = re.match(r"^(#{1,6})\s+(.*)$", ligne)
        if m:
            titre = m.group(2).upper()
            lignes.append(titre)
            lignes.append("-" * len(titre))
        else:
            lignes.append(ligne)
    return "\n".join(lignes)


# ----------------------------------------------------------------------
# Word (.docx)
# ----------------------------------------------------------------------
def _vers_docx(texte: str, titre: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if titre:
        doc.add_heading(titre, level=0)

    for ligne in texte.split("\n"):
        brute = ligne.rstrip()
        if not brute.strip():
            continue
        # Titres Markdown (# ## ###).
        m = re.match(r"^(#{1,6})\s+(.*)$", brute)
        if m:
            niveau = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=niveau)
            continue
        # Listes a puces.
        m = re.match(r"^\s*[-*]\s+(.*)$", brute)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _ajoute_runs_docx(p, m.group(1), Pt)
            continue
        # Listes numerotees.
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", brute)
        if m:
            p = doc.add_paragraph(style="List Number")
            _ajoute_runs_docx(p, m.group(1), Pt)
            continue
        # Paragraphe normal.
        p = doc.add_paragraph()
        _ajoute_runs_docx(p, brute, Pt)

    tampon = io.BytesIO()
    doc.save(tampon)
    return tampon.getvalue()


def _ajoute_runs_docx(paragraphe, texte, Pt):
    for morceau, gras in _segments_gras(texte):
        run = paragraphe.add_run(morceau)
        run.bold = gras
        run.font.size = Pt(11)


# ----------------------------------------------------------------------
# PDF (reportlab)
# ----------------------------------------------------------------------
def _echappe_pdf(texte: str) -> str:
    t = texte.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)


def _vers_pdf(texte: str, titre: str) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    corps = ParagraphStyle(
        "Corps", parent=styles["Normal"], fontSize=11, leading=16, alignment=TA_LEFT
    )

    tampon = io.BytesIO()
    doc = SimpleDocTemplate(
        tampon,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        title=titre or "Document",
    )

    flux = []
    if titre:
        flux.append(Paragraph(_echappe_pdf(titre), styles["Title"]))
        flux.append(Spacer(1, 0.4 * cm))

    puces = []  # accumulateur pour regrouper les puces consecutives

    def vider_puces():
        if puces:
            flux.append(
                ListFlowable(
                    [ListItem(Paragraph(p, corps)) for p in puces],
                    bulletType="bullet",
                    leftIndent=18,
                )
            )
            flux.append(Spacer(1, 0.2 * cm))
            puces.clear()

    for ligne in texte.split("\n"):
        brute = ligne.rstrip()
        if not brute.strip():
            vider_puces()
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", brute)
        if m:
            vider_puces()
            niveau = min(len(m.group(1)), 3)
            style = styles[f"Heading{niveau}"]
            flux.append(Paragraph(_echappe_pdf(m.group(2).strip()), style))
            continue
        m = re.match(r"^\s*[-*]\s+(.*)$", brute)
        if m:
            puces.append(_echappe_pdf(m.group(1)))
            continue
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", brute)
        if m:
            puces.append(_echappe_pdf(m.group(1)))
            continue
        vider_puces()
        flux.append(Paragraph(_echappe_pdf(brute), corps))
        flux.append(Spacer(1, 0.15 * cm))

    vider_puces()
    if not flux:
        flux.append(Paragraph(" ", corps))
    doc.build(flux)
    return tampon.getvalue()


# ----------------------------------------------------------------------
# CV professionnel (PDF / Word)
# ----------------------------------------------------------------------
_COULEUR_CV = "#1a4d8c"
_COULEUR_CV_CLAIR = "#e8f0fa"


def _lignes_cv(texte: str) -> list[tuple[str, str]]:
    """Parse le Markdown CV en (type, contenu). type: h1|h2|h3|p|li|contact."""
    resultat: list[tuple[str, str]] = []
    for ligne in texte.split("\n"):
        brute = ligne.rstrip()
        if not brute.strip():
            continue
        m = re.match(r"^#\s+(.*)$", brute)
        if m:
            resultat.append(("h1", m.group(1).strip()))
            continue
        m = re.match(r"^##\s+(.*)$", brute)
        if m:
            resultat.append(("h2", m.group(1).strip()))
            continue
        m = re.match(r"^###\s+(.*)$", brute)
        if m:
            resultat.append(("h3", m.group(1).strip()))
            continue
        m = re.match(r"^\s*[-*]\s+(.*)$", brute)
        if m:
            resultat.append(("li", m.group(1).strip()))
            continue
        if resultat and resultat[-1][0] == "h1" and "|" in brute:
            resultat.append(("contact", brute.strip()))
            continue
        resultat.append(("p", brute.strip()))
    return resultat


def _vers_cv_pdf(texte: str, titre: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    nom_style = ParagraphStyle(
        "NomCV", parent=styles["Title"], fontSize=22, textColor=colors.white,
        spaceAfter=4, leading=26,
    )
    contact_style = ParagraphStyle(
        "ContactCV", parent=styles["Normal"], fontSize=9, textColor=colors.white,
        spaceAfter=0, leading=12,
    )
    section_style = ParagraphStyle(
        "SectionCV", parent=styles["Heading2"], fontSize=12, textColor=colors.HexColor(_COULEUR_CV),
        spaceBefore=10, spaceAfter=4, leading=14,
    )
    poste_style = ParagraphStyle(
        "PosteCV", parent=styles["Heading3"], fontSize=11, textColor=colors.black,
        spaceBefore=6, spaceAfter=2, leading=13,
    )
    corps = ParagraphStyle(
        "CorpsCV", parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_LEFT,
    )
    puce = ParagraphStyle(
        "PuceCV", parent=corps, leftIndent=14, bulletIndent=6, spaceBefore=1,
    )

    tampon = io.BytesIO()
    doc = SimpleDocTemplate(
        tampon, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm, title=titre or "CV",
    )
    flux: list = []
    nom = titre or "Curriculum Vitae"
    contact = ""
    lignes = _lignes_cv(texte)
    if lignes and lignes[0][0] == "h1":
        nom = lignes[0][1]
        lignes = lignes[1:]
    for typ, cont in lignes:
        if typ == "contact":
            contact = cont
            break

    entete_data = [[Paragraph(_echappe_pdf(nom), nom_style)]]
    if contact:
        entete_data.append([Paragraph(_echappe_pdf(contact), contact_style)])
    entete = Table(entete_data, colWidths=[16 * cm])
    entete.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(_COULEUR_CV)),
        ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING", (0, 0), (-1, 0), 12),
        ("BOTTOMPADDING", (0, -1), (-1, -1), 12),
    ]))
    flux.append(entete)
    flux.append(Spacer(1, 0.5 * cm))

    for typ, cont in lignes:
        if typ == "contact":
            continue
        esc = _echappe_pdf(cont)
        if typ == "h1":
            flux.append(Paragraph(esc, nom_style))
        elif typ == "h2":
            flux.append(Spacer(1, 0.15 * cm))
            flux.append(Paragraph(esc, section_style))
            flux.append(Spacer(1, 0.05 * cm))
        elif typ == "h3":
            flux.append(Paragraph(esc, poste_style))
        elif typ == "li":
            flux.append(Paragraph(f"• {esc}", puce))
        else:
            flux.append(Paragraph(esc, corps))
            flux.append(Spacer(1, 0.08 * cm))

    if len(flux) <= 2:
        flux.append(Paragraph(_echappe_pdf(_markdown_simple(texte)), corps))
    doc.build(flux)
    return tampon.getvalue()


def _vers_cv_docx(texte: str, titre: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    lignes = _lignes_cv(texte)
    nom = titre or "Curriculum Vitae"
    if lignes and lignes[0][0] == "h1":
        nom = lignes[0][1]
        lignes = lignes[1:]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(nom)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(26, 77, 140)

    for typ, cont in lignes:
        if typ == "contact":
            p = doc.add_paragraph(cont)
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(80, 80, 80)
        elif typ == "h2":
            doc.add_paragraph()
            p = doc.add_heading(cont, level=2)
            for r in p.runs:
                r.font.color.rgb = RGBColor(26, 77, 140)
        elif typ == "h3":
            p = doc.add_paragraph()
            run = p.add_run(cont)
            run.bold = True
            run.font.size = Pt(11)
        elif typ == "li":
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(cont)
            run.font.size = Pt(10)
        elif typ == "h1":
            pass
        else:
            p = doc.add_paragraph(cont)
            for r in p.runs:
                r.font.size = Pt(10)

    tampon = io.BytesIO()
    doc.save(tampon)
    return tampon.getvalue()


# ----------------------------------------------------------------------
# Point d'entree
# ----------------------------------------------------------------------
def _retire_titre_redondant(texte: str, titre: str) -> str:
    """Si le corps commence par un titre Markdown identique au titre du document,
    on le retire pour eviter de l'afficher deux fois (docx/pdf gerent le titre)."""
    if not titre:
        return texte
    lignes = texte.split("\n")
    for i, ligne in enumerate(lignes):
        if not ligne.strip():
            continue
        m = re.match(r"^#{1,6}\s+(.*)$", ligne.strip())
        if m and m.group(1).strip().lower() == titre.strip().lower():
            return "\n".join(lignes[i + 1:]).lstrip("\n")
        break
    return texte


def generer(texte: str, fmt: str, titre: str = "", type_doc: str = "generic") -> bytes:
    """Transforme un contenu Markdown en fichier du format demande."""
    fmt = (fmt or "txt").lower()
    type_doc = (type_doc or "generic").lower()
    corps = _retire_titre_redondant(texte, titre)
    if fmt == "md":
        return texte.encode("utf-8")
    if fmt == "txt":
        return _markdown_simple(texte).encode("utf-8")
    if type_doc == "cv":
        if fmt == "docx":
            return _vers_cv_docx(corps, titre)
        if fmt == "pdf":
            return _vers_cv_pdf(corps, titre)
    if fmt == "docx":
        return _vers_docx(corps, titre)
    if fmt == "pdf":
        return _vers_pdf(corps, titre)
    raise ValueError(f"format inconnu : {fmt}")
